"""
Microsoft Word (.docx) file loader
"""
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx not installed. DOCX support will not work.")
    print("Install with: pip install python-docx")


def load_docx(file_path: str) -> str:
    """
    Extract text from DOCX file
    
    Args:
        file_path: Path to .docx file
    
    Returns:
        Extracted text from document
    """
    if not HAS_DOCX:
        return "Error: python-docx not installed. Please install python-docx."
    
    try:
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except Exception as e:
        print(f"Error while reading DOCX: {e}")
        return f"Error reading DOCX file: {str(e)}"
